"""
Django management command to load client DOCX files into knowledge_base with embeddings.
Run: python manage.py load_docx

Ported from Flask scripts/load_docx_to_knowledge.py
"""
import re
import json
from pathlib import Path
from django.core.management.base import BaseCommand
from django.conf import settings
from docx import Document


# State name from filename: "Respostas Arizona.docx" -> "Arizona"
RESPOSTAS_PREFIX = "Respostas "
QUESTIONS_FILE = "Lista de Perguntas - IA.docx"


def get_paragraphs(doc_path: Path) -> list:
    """Extract non-empty paragraph texts from a DOCX."""
    doc = Document(doc_path)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def get_questions(doc_path: Path) -> list:
    """Extract questions from Lista de Perguntas doc (paragraphs or table)."""
    doc = Document(doc_path)
    questions = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            questions.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t and t not in questions:
                    questions.append(t)
    return questions


def parse_qa_pairs(paragraphs: list) -> list:
    """If doc has 'P: ... R: ...' or 'Q: ... A: ...' style, return [(q, a), ...]."""
    pairs = []
    i = 0
    while i < len(paragraphs):
        line = paragraphs[i]
        q_match = re.match(r"^(?:P|Pergunta|Q|Question)\s*[:\d.]*\s*(.+)", line, re.I) or re.match(r"^(.+)\?$", line)
        if q_match and i + 1 < len(paragraphs):
            pairs.append((q_match.group(1).strip(), paragraphs[i + 1].strip()))
            i += 2
            continue
        i += 1
    return pairs


def parse_question_answer_paragraphs(paragraphs: list) -> list:
    """Pair paragraphs: if a paragraph ends with '?', use it as question and next paragraph as answer."""
    pairs = []
    i = 0
    while i < len(paragraphs):
        q = paragraphs[i].strip()
        if not q:
            i += 1
            continue
        if q.rstrip().endswith("?") and i + 1 < len(paragraphs):
            a = paragraphs[i + 1].strip()
            if a:
                pairs.append((q, a))
                i += 2
                continue
        pairs.append((q, q))
        i += 1
    return pairs


def _is_question_line(text: str) -> bool:
    """True if line looks like a question (ends with ?) or a numbered question."""
    t = text.strip()
    if not t:
        return False
    if t.rstrip().endswith("?"):
        return True
    if re.match(r"^\d+\.\s+", t):
        return True
    return False


def parse_question_answer_multiparagraph(paragraphs: list) -> list:
    """
    Pair question with full answer: question = line ending with ? or starting with "N. ".
    Answer = all following paragraphs until the next question line.
    """
    pairs = []
    i = 0
    while i < len(paragraphs):
        line = paragraphs[i].strip()
        if not line:
            i += 1
            continue
        if not _is_question_line(line):
            i += 1
            continue
        q = line
        answer_lines = []
        j = i + 1
        while j < len(paragraphs):
            next_line = paragraphs[j].strip()
            if not next_line:
                j += 1
                continue
            if _is_question_line(next_line):
                break
            answer_lines.append(next_line)
            j += 1
        full_answer = "\n".join(answer_lines).strip() if answer_lines else q
        pairs.append((q, full_answer))
        i = j if answer_lines else i + 1
    return pairs


class Command(BaseCommand):
    help = "Load client DOCX files into knowledge_base with embeddings"

    def handle(self, *args, **options):
        from chat.models import KnowledgeBase
        from chat.services.knowledge_service import get_embedding

        _backend_dir = Path(settings.BASE_DIR)

        # Try multiple locations
        candidates = [
            Path(settings.DOCX_DATA_DIR),
            _backend_dir.parent,   # project root
            _backend_dir,          # django_backend folder
            _backend_dir / "data",
            _backend_dir.parent / "documents",
        ]
        data_dir = None
        for d in candidates:
            if d.is_dir():
                if (d / QUESTIONS_FILE).exists() or list(d.glob("Respostas *.docx")):
                    data_dir = d
                    break
        if data_dir is None:
            data_dir = _backend_dir.parent if _backend_dir.parent.is_dir() else _backend_dir
        self.stdout.write(f"Using data dir: {data_dir}")

        questions_file = data_dir / QUESTIONS_FILE
        questions = []
        if questions_file.exists():
            questions = get_questions(questions_file)
            self.stdout.write(f"Loaded {len(questions)} questions from {QUESTIONS_FILE}")
        else:
            self.stdout.write(f"Questions file not found: {questions_file}")

        # Respostas files: state -> list of answers
        state_answers = {}
        for path in sorted(data_dir.glob("*.docx")):
            name = path.stem
            if name.startswith(RESPOSTAS_PREFIX):
                state = name[len(RESPOSTAS_PREFIX):].strip()
                paras = get_paragraphs(path)
                if not paras:
                    continue
                if questions and len(paras) == len(questions):
                    state_answers[state] = list(zip(questions, paras))
                else:
                    pairs = parse_qa_pairs(paras)
                    if pairs:
                        state_answers[state] = pairs
                    else:
                        pairs = parse_question_answer_multiparagraph(paras)
                        if pairs:
                            state_answers[state] = pairs
                        else:
                            pairs = parse_question_answer_paragraphs(paras)
                            if pairs:
                                state_answers[state] = pairs
                            else:
                                state_answers[state] = [(p, p) for p in paras]
                self.stdout.write(f"  {state}: {len(state_answers[state])} entries")

        if not state_answers and not questions:
            self.stdout.write("No DOCX data found. Add Lista de Perguntas - IA.docx and Respostas <State>.docx files.")
            return

        # Build (state, question, answer) list
        entries = []
        for state, qa_list in state_answers.items():
            for q, a in qa_list:
                if q and a:
                    entries.append((state, q, a))
        for q in questions:
            if q and not any(e[1] == q for e in entries):
                entries.append((None, q, "See state-specific answers or ask for your state."))

        if not settings.OPENAI_API_KEY:
            self.stdout.write("OPENAI_API_KEY not set. Embeddings will be empty; semantic search will not work until you add key and re-run.")

        existing = KnowledgeBase.objects.count()
        if existing > 0:
            KnowledgeBase.objects.all().delete()
            self.stdout.write(f"Cleared {existing} existing knowledge_base rows.")

        added = 0
        for state, question, answer in entries:
            emb = get_embedding(question)
            emb_json = json.dumps(emb) if emb else None
            KnowledgeBase.objects.create(
                state=state,
                question=question,
                answer=answer,
                embedding_json=emb_json,
            )
            added += 1
            if added % 10 == 0:
                self.stdout.write(f"  Added {added}/{len(entries)}...")

        self.stdout.write(self.style.SUCCESS(f"Done. Inserted {added} rows into knowledge_base."))
