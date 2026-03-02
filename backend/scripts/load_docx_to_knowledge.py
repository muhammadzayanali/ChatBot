"""
Load client DOCX files into knowledge_base with embeddings.
Run from project root: python -m backend.scripts.load_docx_to_knowledge
Or from backend: python scripts/load_docx_to_knowledge.py (after adding parent to path)

Expected files (in project root or DOCX_DATA_DIR):
- Lista de Perguntas - IA.docx  (list of questions)
- Respostas Arizona.docx, Respostas NY.docx, Respostas Texas.docx, Respostas Florida.docx,
  Respostas Colorado.docx, Respostas Illinois.docx, Respostas California.docx, Respostas Pennsylvania.docx

Each Respostas file: one answer per paragraph (same order as questions) or Q/A pairs.
"""
import sys
import re
from pathlib import Path

# Allow running from project root or backend
_backend_dir = Path(__file__).resolve().parent.parent
if str(_backend_dir) not in sys.path:
    sys.path.insert(0, str(_backend_dir))

# Load .env from backend so OPENAI_API_KEY is set for embeddings
try:
    from dotenv import load_dotenv
    load_dotenv(_backend_dir / ".env")
except ImportError:
    pass

from docx import Document
from config import DOCX_DATA_DIR, OPENAI_API_KEY
from database.models import init_db, SessionLocal, KnowledgeBase
from services.knowledge_service import get_embedding

# State name from filename: "Respostas Arizona.docx" -> "Arizona"
RESPOSTAS_PREFIX = "Respostas "
QUESTIONS_FILE = "Lista de Perguntas - IA.docx"


def get_paragraphs(doc_path: Path) -> list:
    """Extract non-empty paragraph texts from a DOCX."""
    doc = Document(doc_path)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def get_questions(doc_path: Path) -> list:
    """Extract questions from Lista de Perguntas doc (paragraphs or table)."""
    doc = Document(doc_path)
    questions = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            questions.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t and t not in questions:
                    questions.append(t)
    return questions


def parse_qa_pairs(paragraphs: list) -> list:
    """If doc has 'P: ... R: ...' or 'Q: ... A: ...' style, return [(q, a), ...]."""
    pairs = []
    i = 0
    while i < len(paragraphs):
        line = paragraphs[i]
        # Portuguese/Spanish/English patterns
        q_match = re.match(r"^(?:P|Pergunta|Q|Question)\s*[:\d.]*\s*(.+)", line, re.I) or re.match(r"^(.+)\?$", line)
        if q_match and i + 1 < len(paragraphs):
            pairs.append((q_match.group(1).strip(), paragraphs[i + 1].strip()))
            i += 2
            continue
        i += 1
    return pairs


def parse_question_answer_paragraphs(paragraphs: list) -> list:
    """Pair paragraphs: if a paragraph ends with '?', use it as question and next paragraph as answer."""
    pairs = []
    i = 0
    while i < len(paragraphs):
        q = paragraphs[i].strip()
        if not q:
            i += 1
            continue
        if q.rstrip().endswith("?") and i + 1 < len(paragraphs):
            a = paragraphs[i + 1].strip()
            if a:
                pairs.append((q, a))
                i += 2
                continue
        # No question mark or no next para: treat as Q&A with same text (answer might be in same para)
        pairs.append((q, q))
        i += 1
    return pairs


def _is_question_line(text: str) -> bool:
    """True if line looks like a question (ends with ?) or a numbered question (e.g. '16. Como...')."""
    t = text.strip()
    if not t:
        return False
    if t.rstrip().endswith("?"):
        return True
    if re.match(r"^\d+\.\s+", t):
        return True
    return False


def parse_question_answer_multiparagraph(paragraphs: list) -> list:
    """
    Pair question with full answer: question = line ending with ? or starting with "N. ".
    Answer = all following paragraphs until the next question line (inclusive of bullets/lists).
    Returns [(question, full_answer_text), ...] with full_answer_text containing newlines.
    """
    pairs = []
    i = 0
    while i < len(paragraphs):
        line = paragraphs[i].strip()
        if not line:
            i += 1
            continue
        if not _is_question_line(line):
            i += 1
            continue
        q = line
        answer_lines = []
        j = i + 1
        while j < len(paragraphs):
            next_line = paragraphs[j].strip()
            if not next_line:
                j += 1
                continue
            if _is_question_line(next_line):
                break
            answer_lines.append(next_line)
            j += 1
        full_answer = "\n".join(answer_lines).strip() if answer_lines else q
        pairs.append((q, full_answer))
        i = j if answer_lines else i + 1
    return pairs


def main():
    # Try multiple locations so we find client DOCX regardless of where they're placed
    candidates = [
        Path(DOCX_DATA_DIR),
        _backend_dir.parent,  # project root
        _backend_dir,         # backend folder
        _backend_dir / "data",
        _backend_dir.parent / "documents",
    ]
    data_dir = None
    for d in candidates:
        if d.is_dir():
            # Check that we have at least one Respostas file or the questions file
            if (d / QUESTIONS_FILE).exists() or list(d.glob("Respostas *.docx")):
                data_dir = d
                break
    if data_dir is None:
        data_dir = _backend_dir.parent if _backend_dir.parent.is_dir() else _backend_dir
    print(f"Using data dir: {data_dir}")

    init_db()
    questions_file = data_dir / QUESTIONS_FILE
    questions = []
    if questions_file.exists():
        questions = get_questions(questions_file)
        print(f"Loaded {len(questions)} questions from {QUESTIONS_FILE}")
    else:
        print(f"Questions file not found: {questions_file}")

    # Respostas files: state -> list of answers (or Q/A pairs)
    state_answers = {}
    for path in sorted(data_dir.glob("*.docx")):
        name = path.stem
        if name.startswith(RESPOSTAS_PREFIX):
            state = name[len(RESPOSTAS_PREFIX):].strip()
            paras = get_paragraphs(path)
            if not paras:
                continue
            # If we have questions and same count, pair by index
            if questions and len(paras) == len(questions):
                state_answers[state] = list(zip(questions, paras))
            else:
                pairs = parse_qa_pairs(paras)
                if pairs:
                    state_answers[state] = pairs
                else:
                    # Prefer full multi-paragraph answers (question? then all bullets/lines until next question)
                    pairs = parse_question_answer_multiparagraph(paras)
                    if pairs:
                        state_answers[state] = pairs
                    else:
                        # Fallback: single question + single answer paragraph
                        pairs = parse_question_answer_paragraphs(paras)
                        if pairs:
                            state_answers[state] = pairs
                        else:
                            state_answers[state] = [(p, p) for p in paras]
            print(f"  {state}: {len(state_answers[state])} entries")

    if not state_answers and not questions:
        print("No DOCX data found. Add Lista de Perguntas - IA.docx and Respostas <State>.docx files.")
        return

    # Build (state, question, answer) list
    entries = []
    for state, qa_list in state_answers.items():
        for q, a in qa_list:
            if q and a:
                entries.append((state, q, a))
    # General entries from questions only (no state)
    for q in questions:
        if q and not any(e[1] == q for e in entries):
            entries.append((None, q, "See state-specific answers or ask for your state."))

    if not OPENAI_API_KEY:
        print("OPENAI_API_KEY not set. Embeddings will be empty; semantic search will not work until you add key and re-run.")
    db = SessionLocal()
    try:
        existing = db.query(KnowledgeBase).count()
        if existing > 0:
            db.query(KnowledgeBase).delete()
            db.commit()
            print(f"Cleared {existing} existing knowledge_base rows.")
        added = 0
        for state, question, answer in entries:
            emb = get_embedding(question)
            emb_json = None
            if emb:
                import json
                emb_json = json.dumps(emb)
            row = KnowledgeBase(
                state=state,
                question=question,
                answer=answer,
                embedding_json=emb_json,
            )
            db.add(row)
            added += 1
            if added % 10 == 0:
                print(f"  Added {added}/{len(entries)}...")
        db.commit()
        print(f"Done. Inserted {added} rows into knowledge_base.")
    finally:
        db.close()


if __name__ == "__main__":
    main()
